# -*- coding: utf-8 -*-
"""将演示文稿和视频文稿 Markdown 转为 Word (.docx)"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_paragraph_with_format(doc, line):
    """处理单行：加粗、普通文本"""
    p = doc.add_paragraph()
    rest = line.strip()
    while rest:
        m = re.match(r'\*\*(.+?)\*\*', rest)
        if m:
            before = rest[:m.start()]
            if before:
                p.add_run(before)
            p.add_run(m.group(1)).bold = True
            rest = rest[m.end():]
        else:
            p.add_run(rest)
            break
    return p

def parse_table(lines, start_i):
    """从 lines[start_i] 开始解析表格，返回 (rows, next_index)"""
    rows = []
    i = start_i
    while i < len(lines):
        line = lines[i]
        if not line.strip() or not line.strip().startswith('|'):
            break
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i

def md_to_docx(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()
        # 标题
        if raw.startswith('# '):
            doc.add_heading(raw[2:].strip(), level=0)
        elif raw.startswith('## '):
            doc.add_heading(raw[3:].strip(), level=1)
        elif raw.startswith('### '):
            doc.add_heading(raw[4:].strip(), level=2)
        # 分隔线
        elif raw.strip() == '---':
            doc.add_paragraph('—' * 20)
        # 引用块
        elif raw.startswith('> '):
            p = doc.add_paragraph()
            p.add_run(raw[2:].strip())
            p.paragraph_format.left_indent = Inches(0.25)
        # 表格
        elif raw.strip().startswith('|'):
            rows, next_i = parse_table(lines, i)
            i = next_i - 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell
                doc.add_paragraph()
        # 无序列表
        elif raw.strip().startswith('- '):
            text = raw.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            # 简单处理加粗
            while '**' in text:
                a, rest = text.split('**', 1)
                if a:
                    p.add_run(a)
                if '**' in rest:
                    b, rest = rest.split('**', 1)
                    p.add_run(b).bold = True
                    text = rest
                else:
                    p.add_run(rest).bold = True
                    text = ''
                    break
            if text:
                p.add_run(text)
        # 有序列表 (1. 2. 等)
        elif re.match(r'^\s*\d+\.\s', raw):
            text = re.sub(r'^\s*\d+\.\s*', '', raw).strip()
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
        # 空行
        elif not raw.strip():
            pass
        # **标签：** 开头的键值
        elif raw.strip().startswith('**') and '：**' in raw or ':**' in raw:
            doc.add_paragraph(raw.strip())
        # 普通段落
        else:
            add_paragraph_with_format(doc, raw)
        i += 1
    doc.save(docx_path)
    print('已生成:', docx_path)

if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    md_to_docx(
        os.path.join(base, '赫兹使用与PC端配合-演示文稿.md'),
        os.path.join(base, '赫兹使用与PC端配合-演示文稿.docx')
    )
    md_to_docx(
        os.path.join(base, '赫兹使用与PC端配合-视频文稿.md'),
        os.path.join(base, '赫兹使用与PC端配合-视频文稿.docx')
    )
